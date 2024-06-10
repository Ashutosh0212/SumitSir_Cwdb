

def get_financial_year():
    # Logic to determine the Indian financial year
    # This is just a sample logic, you can adjust it as per your needs
    from datetime import datetime
    today = datetime.today()
    if today.month >= 4:
        return f'{today.year}-{today.year + 1}'
    else:
        return f'{today.year - 1}-{today.year}'

from cwdb_admin.models import Proposal,SanctionLetter,InspectionReport,FundDistribution
# from datetime import datetime
from datetime import datetime

QUARTER_CHOICES = [
    ('Q1', 'Quarter 1 (April-June)'),
    ('Q2', 'Quarter 2 (July-September)'),
    ('Q3', 'Quarter 3 (October-December)'),
    ('Q4', 'Quarter 4 (January-March)'),
]

def get_current_quarter():
    now = datetime.now()
    mon=now.month
    if mon>=4 and mon<=6:
        quarter=1
    elif mon>=7 and mon<=9:
        quarter=2
    elif mon>=10 and mon<=12:
        quarter=3
    else:
        quarter=4
    current_quarter = f'Q{quarter}'
    return current_quarter

# print(get_current_quarter())

from .forms import CustomUserCreationForm


# class SignUp(generic.CreateView):
#     form_class = CustomUserCreationForm
#     success_url = reverse_lazy("login")
#     template_name = "signup.html"ss

from django.http import HttpResponse  
from django.shortcuts import render, redirect  
from django.contrib.auth import login, authenticate  
from django.contrib.sites.shortcuts import get_current_site  
from django.utils.encoding import force_bytes, force_str  
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode  
from django.template.loader import render_to_string  
from authapp.tokens import account_activation_token  
from django.contrib.auth.models import User  
from django.core.mail import EmailMessage,send_mail
from src.settings import EMAIL_HOST_USER
def signup(request):  
    if request.method == 'POST':  
        form = CustomUserCreationForm(request.POST)  
        if form.is_valid():  
            # save form in the memory not in database  
            user = form.save(commit=False)  
            user.is_active = False  
            user.save()  
            # to get the domain of the current site  
            current_site = get_current_site(request)  
            mail_subject = 'Activation link has been sent to your email id'  
            message = render_to_string('registration/acc_active_email.html', {  
                'user': user,  
                'domain': current_site.domain,  
                'uid':urlsafe_base64_encode(force_bytes(user.pk)),  
                'token':account_activation_token.make_token(user),  
            })  
            to_email = form.cleaned_data.get('email')  
            send_mail(mail_subject, message, EMAIL_HOST_USER, [to_email])
            # print(EMAIL_HOST_USER)
            # print(to_email)
            # print("Email Content:")
            # print(message) 
              
            # print(email)
            return render(request, 'registration/registration_confirmation.html')  
    else:  
        form = CustomUserCreationForm()  
    return render(request, 'registration/signup.html', {'form': form})

from django.contrib.auth import authenticate, login
from django.shortcuts import render, redirect


from django.contrib.auth.views import LoginView
from django.shortcuts import redirect



from django.contrib.auth import get_user_model

def activate(request, uidb64, token):  
    User = get_user_model()  
    try:  
        uid = force_str(urlsafe_base64_decode(uidb64))  
        user = User.objects.get(pk=uid)  
    except(TypeError, ValueError, OverflowError, User.DoesNotExist):  
        user = None  
    if user is not None and account_activation_token.check_token(user, token):  
        user.is_active = True  
        user.save()  
        return render(request, 'registration/activation_confirmation.html')  
    else:
        if user is not None:
            user.delete()  #
        return HttpResponse('Activation link is invalid!')  
    
# views.py


#profile
from django.shortcuts import render, redirect
from django.views import View
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from .models import CustomUser
from .forms import CustomUserChangeForm

@method_decorator(login_required, name='dispatch')
class UserProfileView(View):
    template_name = 'registration/user_profile.html'

    def get(self, request, *args, **kwargs):
        user = request.user
        return render(request, self.template_name, {'user': user})

@method_decorator(login_required, name='dispatch')
class EditProfileView(View):
    template_name = 'registration/edit_profile.html'

    def get(self, request, *args, **kwargs):
        user = request.user
        form = CustomUserChangeForm(instance=user)
        return render(request, self.template_name, {'user': user, 'form': form})

    def post(self, request, *args, **kwargs):
        user = request.user
        form = CustomUserChangeForm(request.POST, instance=user)
        if form.is_valid():
            form.save()
            return redirect('authapp:user_profile')  # Redirect to the user profile page after a successful update
        return render(request, self.template_name, {'user': user, 'form': form})





#admin
# views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.core.mail import send_mail
from django.contrib.auth.decorators import user_passes_test
from .forms import ProposalApprovalForm




from django.core.mail import EmailMessage
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from .models import Notification
from django.utils import timezone

def send_status_change_notification(user_email, project_id, new_status, sanction_letter=None):
    attachment_name = None
    subject = 'Your Proposal Status has Changed'
    context = {'project_id': project_id, 'new_status': new_status}
    message_html = render_to_string('email/notification_template.html', context)
    message_plain = strip_tags(message_html)
    
    notification_message_html=render_to_string('email/notification_template_dashboard.html', context)
    notification_message_plain=strip_tags(notification_message_html)
    # Save the notification to the model
    notification = Notification.objects.create(
        user=user_email,  # Assuming user_email is a User object
        message=notification_message_plain,  # You can customize this based on your needs
        created_at=timezone.now(),
    )
    
    email = EmailMessage(subject, message_html, from_email=EMAIL_HOST_USER, to=[user_email])
    email.content_subtype = 'html'  # Set the content type to HTML

    if new_status == "Approved" and sanction_letter:
        attachment_name = 'sanction_letter.pdf'
        email.attach(attachment_name, sanction_letter.read(), 'application/pdf')
    
    # Save the attachment to the notification model
   
    notification.save()

    email.send()


#show notifications view
# views.py
from django.shortcuts import render
from .models import Notification
from django.contrib.auth.decorators import login_required


@login_required
def show_all_notifications(request):
    # Retrieve all notifications for the logged-in user
    all_notifications = Notification.objects.filter(user=request.user).order_by('-created_at')

    return render(request, 'all_notifications.html', {'all_notifications': all_notifications})

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
# from .models import FormSubmission
from django.utils.crypto import get_random_string

@login_required
def dashboard(request):
    #top three notificatons
    notifications = Notification.objects.filter(user=request.user).order_by('-created_at')
    first_three_notifications = notifications[:3]
     # Count of all proposals submitted by the logged-in user
    submitted_proposals_count = Proposal.objects.filter(user=request.user).count()
    # Count of pending proposals
    pending_proposals_count = Proposal.objects.filter(user=request.user, status='Pending').count()
     # Count of approved proposals
    approved_proposals_count = Proposal.objects.filter(user=request.user, status='Approved').count()
 # Count of proposals to be resubmitted
    resubmit_proposals_count = Proposal.objects.filter(user=request.user, status='Resubmitted').count()
    #count proposals rejected
    rejected_proposals_count=Proposal.objects.filter(user=request.user, status='Rejected').count()

    return render(request, 'main/dashboard.html',{'submitted_proposals_count': submitted_proposals_count,
        'pending_proposals_count': pending_proposals_count,
        'approved_proposals_count': approved_proposals_count,
        'rejected_proposals_count': rejected_proposals_count,
        'new_notifications': first_three_notifications})

SCHEME_CHOICES = [
        ('WMS', 'WMS'),
        ('WPS', 'WPS'),
        ('HRD', 'HRD'),
        ('PWDS', 'PWDS'),
    ]

from django.shortcuts import render
from .models import FundDistribution,Proposal,BeneficiaryData,ExpenditureData
from django.db.models import Sum 
from .forms import quarterly_schemes_form,allocation_form

def index(request):
    current_financial_year = get_financial_year()

    # Fetch data for the current financial year
    try:
        fund_distribution_data = FundDistribution.objects.get(financial_year=current_financial_year)
    except FundDistribution.DoesNotExist:
        fund_distribution_data = None  # Handle the case where data for the current year does not exist
        
      # Count the number of approved and completed projects
    approved_projects_count = Proposal.objects.filter(status='Approved').count()
    completed_projects_count = Proposal.objects.filter(status='Completed').count()
    total_projects_count = approved_projects_count + completed_projects_count
 # Sum up the num_beneficiaries of every row in the BeneficiaryData table
    total_beneficiaries = BeneficiaryData.objects.aggregate(Sum('num_beneficiaries'))['num_beneficiaries__sum']
    
    #fetch data according to form 
    #quarter expenditure and sanctioned data
    form = quarterly_schemes_form(request.GET or None)
    form1=allocation_form(request.GET or None)
    # print(f"Current Financial Year: {current_financial_year}")
    if form1.is_valid():
        year = form1.cleaned_data.get('financial_year')
        try:
            fund_data = FundDistribution.objects.get(financial_year=year)
        except FundDistribution.DoesNotExist:
            fund_data = None  # Handle the case where data for the current year does not exist
    else:
        fund_data=fund_distribution_data
        
    schemes_data = {}
    if form.is_valid():
        year = form.cleaned_data.get('financial_year')
        budget_type=form.cleaned_data.get('select_type')
    else:
        year=current_financial_year
        budget_type='Fund Sanctioned'
        
    if budget_type == 'Expenditure':
        schemes_data = {scheme[0]: {quarter: 0 for quarter in ['Q1', 'Q2', 'Q3', 'Q4']} for scheme in SCHEME_CHOICES}
        for quarter in schemes_data['WMS']:
            quarterly_sums = ExpenditureData.get_quarterly_sums(year, quarter)
            for scheme, value in quarterly_sums.items():
                schemes_data[scheme][quarter] = value
    elif budget_type == 'Fund Sanctioned':
        schemes_data = {scheme[0]: {quarter: 0 for quarter in ['Q1', 'Q2', 'Q3', 'Q4']} for scheme in SCHEME_CHOICES}
        for quarter in schemes_data['WMS']:
            quarterly_sums_allocated = ExpenditureData.get_quarterly_sums_allocated(year, quarter)
            for scheme, value in quarterly_sums_allocated.items():
                schemes_data[scheme][quarter] = value

        # Additional data for more charts
    quarterlyFunds = {
        'quarters': ['Q1', 'Q2', 'Q3', 'Q4'],
        'scheme1Funds': [schemes_data.get('WMS', {}).get('Q1', 0), schemes_data.get('WMS', {}).get('Q2', 0), schemes_data.get('WMS', {}).get('Q3', 0), schemes_data.get('WMS', {}).get('Q4', 0)],
        'scheme2Funds': [schemes_data.get('WPS', {}).get('Q1', 0), schemes_data.get('WPS', {}).get('Q2', 0), schemes_data.get('WPS', {}).get('Q3', 0), schemes_data.get('WPS', {}).get('Q4', 0)],
        'scheme3Funds': [schemes_data.get('HRD', {}).get('Q1', 0), schemes_data.get('HRD', {}).get('Q2', 0), schemes_data.get('HRD', {}).get('Q3', 0), schemes_data.get('HRD', {}).get('Q4', 0)],
        'scheme4Funds': [schemes_data.get('PWDS', {}).get('Q1', 0), schemes_data.get('PWDS', {}).get('Q2', 0), schemes_data.get('PWDS', {}).get('Q3', 0), schemes_data.get('PWDS', {}).get('Q4', 0)],
    }
    
    # print(quarterlyFunds)
    context={'fund_distribution_data': fund_distribution_data, 'current_financial_year': current_financial_year, 'total_projects_count': total_projects_count, 'total_beneficiaries': total_beneficiaries,'quarterlyFunds': quarterlyFunds,'quarter_form':form,'schemes_data': schemes_data
                                               ,'fund_data_allocated':fund_data,'allocation_form':form1}
    
   
    return render(request, 'main/index.html',context )

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .forms import FileUploadForm
from .models import Proposal
import random
import string

# views.py
from django.shortcuts import render, redirect
from .models import Proposal
from django.contrib.auth.decorators import login_required
import uuid

STATE_CHOICES = [
    ('Andhra Pradesh', 'Andhra Pradesh'),
    ('Arunachal Pradesh', 'Arunachal Pradesh'),
    ('Assam', 'Assam'),
    ('Bihar', 'Bihar'),
    ('Chhattisgarh', 'Chhattisgarh'),
    ('Goa', 'Goa'),
    ('Gujarat', 'Gujarat'),
    ('Haryana', 'Haryana'),
    ('Himachal Pradesh', 'Himachal Pradesh'),
    ('Jharkhand', 'Jharkhand'),
    ('Karnataka', 'Karnataka'),
    ('Kerala', 'Kerala'),
    ('Madhya Pradesh', 'Madhya Pradesh'),
    ('Maharashtra', 'Maharashtra'),
    ('Manipur', 'Manipur'),
    ('Meghalaya', 'Meghalaya'),
    ('Mizoram', 'Mizoram'),
    ('Nagaland', 'Nagaland'),
    ('Odisha', 'Odisha'),
    ('Punjab', 'Punjab'),
    ('Rajasthan', 'Rajasthan'),
    ('Sikkim', 'Sikkim'),
    ('Tamil Nadu', 'Tamil Nadu'),
    ('Telangana', 'Telangana'),
    ('Tripura', 'Tripura'),
    ('Uttarakhand', 'Uttarakhand'),
    ('Uttar Pradesh', 'Uttar Pradesh'),
    ('West Bengal', 'West Bengal'),
    ('Andaman and Nicobar Islands', 'Andaman and Nicobar Islands'),
    ('Chandigarh', 'Chandigarh'),
    ('Dadra and Nagar Haveli and Daman and Diu', 'Dadra and Nagar Haveli and Daman and Diu'),
    ('Delhi', 'Delhi'),
    ('Jammu and Kashmir', 'Jammu and Kashmir'),
    ('Ladakh', 'Ladakh'),
    ('Lakshadweep', 'Lakshadweep'),
    ('Puducherry', 'Puducherry'),
]

from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse

@login_required
def submit_proposal(request):
    if request.method == 'POST':
        
        agency_address = request.POST.get('agencyAddress')
        implementingAgencyState=request.POST.get('implementingAgencyState')
        project_scheme = request.POST.get('projectScheme')
        scheme_component = request.POST.get('schemeComponent')
        applicant_nature = request.POST.get('applicantNature')
        other_nature = request.POST.get('otherNature')
        brief_of_agency = request.POST.get('agencyActivities')
        objectives_of_project = request.POST.get('Objectives')
        brief_of_project = request.POST.get('Brief')
        justification_of_project = request.POST.get('justification')
        methodology_of_project = request.POST.get('Methodology')
        # expected_outcome = request.POST.get('Outcome')
        scenario_change = request.POST.get('change-scenario')
        # beneficiaries = request.POST.get('beneficiaries')
        mode_of_selection = request.POST.get('mode')
        # component_wise_cost = request.POST.get('project Cost')
        # total_duration = request.POST.get('Total duration')
        location_of_project = request.POST.get('projectLocation')
        associated_agency = request.POST.get('associatedAgency')
        bank_details = request.POST.get('bankAccountDetails')
        nodal_officer_details = request.POST.get('nodalOfficerInfo')
        other_info = request.POST.get('otherInfo')
        quarters = request.POST.get('quarters')
        # quarterly_goal_classes = request.POST.getlist('goal_texts')
        
        # Handle file uploads
        expected_outcome_file = request.POST.get('outcomeFile')
        beneficiaries_file = request.FILES.get('beneficiariesFile')
        component_wise_cost_file = request.FILES.get('projectCostFile')
        total_duration_file = request.FILES.get('durationFile')
        project_report_file = request.FILES.get('projectReport')
        covering_letter_file = request.FILES.get('coveringLetter')

        # Handle dynamic goals
        goals_data = []
        for i in range(1, int(quarters) + 1):
            goal_key = f'goal_texts_quarter:{i}[]'
            goal_texts = request.POST.getlist(goal_key)
            quarter_goals = dict()
            for index, text in enumerate(goal_texts):
                # Initialize each goal as "No" (not submitted)
                quarter_goals[f'goal_{index + 1}'] = {'text': text, 'completed': 0}
            goals_data.append(quarter_goals)


        # Create and save the Proposal object
        proposal = Proposal(
            user=request.user,
            name_and_address=agency_address,
            implementingAgencyState=implementingAgencyState,
            project_scheme=project_scheme,
            scheme_component=scheme_component,
            nature_of_applicant=applicant_nature,
            other_nature=other_nature,
            brief_of_agency=brief_of_agency,
            objectives_of_project=objectives_of_project,
            brief_of_project=brief_of_project,
            justification_of_project=justification_of_project,
            methodology_of_project=methodology_of_project,
            # expected_outcome=expected_outcome,
            scenario_change=scenario_change,
            # beneficiaries=beneficiaries,
            mode_of_selection=mode_of_selection,
            # component_wise_cost=component_wise_cost,
            # total_duration=total_duration,
            location_of_project=location_of_project,
            associated_agency=associated_agency,
            bank_details=bank_details,
            nodal_officer_details=nodal_officer_details,
            other_info=other_info,
            expected_outcome=expected_outcome_file,
            beneficiaries=beneficiaries_file,
            component_wise_cost=component_wise_cost_file,
            component_wise_duration=total_duration_file,
            project_report=project_report_file,
            covering_letter=covering_letter_file,
            total_duration=quarters ,
            goals = goals_data 
        )

        # Generate a unique proposal_id
    
        proposal.unique_id = generate_unique_id(project_scheme)
        # print(proposal)
        proposal.save()

        return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    
    # Pass STATE_CHOICES to the context
    context = {
        'STATE_CHOICES': STATE_CHOICES,
    }

    return render(request, 'proposal/submit_proposal.html', context)


    

import uuid
from datetime import datetime

def generate_unique_id(scheme):
    # Get the current date and year
    current_date = datetime.now()
    year = current_date.year

    # Generate a unique ID using scheme, date, and year
    unique_id = f"{scheme}-{current_date.strftime('%Y%m%d')}-{uuid.uuid4().hex[:4]}"
    return unique_id

@login_required
def proposal_status(request):
    # Filter proposals for the logged-in user
    user_proposals = Proposal.objects.filter(user=request.user)

    return render(request, 'proposal/proposal_status.html', {'proposals': user_proposals})

@login_required
def proposal_detail(request, proposal_id):
    proposal = Proposal.objects.get(id=proposal_id)
    return render(request, 'proposal/proposal_detail.html', {'proposal': proposal})

@login_required
def proposal_list(request):
    proposals = Proposal.objects.filter(user=request.user)
    return render(request, 'proposal/proposal_list.html', {'proposals': proposals})

@login_required
def admin_proposal_list(request):
    proposals = Proposal.objects.all()
    return render(request, 'proposal/admin_proposal_list.html', {'proposals': proposals})

@login_required
def admin_proposal_detail(request, proposal_id):
    proposal = Proposal.objects.get(id=proposal_id)
    if request.method == 'POST':
        status = request.POST.get('status')
        proposal.status = status
        proposal.save()
        return redirect('admin_proposal_detail', proposal_id=proposal.id)
    return render(request, 'proposal/admin_proposal_detail.html', {'proposal': proposal})


from django.shortcuts import render, redirect
from .models import WMS_RevolvingFund
from .forms import WMSRevolvingFundForm
from django.http import HttpResponse



from django.core.mail import EmailMessage
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import user_passes_test
from .models import Proposal, SanctionLetter, InspectionReport, Notification
from .forms import SanctionLetterForm, InspectionReportForm

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def submit_installment_sanction_letter(request, proposal_id):
    proposal = get_object_or_404(Proposal, unique_id=proposal_id)

    if request.method == 'POST':
        form = SanctionLetterForm(request.POST, request.FILES)
        if form.is_valid():
            # Save the form to create a new SanctionLetter instance
            sanction_letter = form.save(commit=False)
            sanction_letter.proposal = proposal
            sanction_letter.save()
            installment_number=sanction_letter.installment_number
            # Send email to user
            send_sanction_letter_email(proposal, sanction_letter,installment_number)

            # Create notification
            create_notification(proposal.user, f"You have received your {sanction_letter.installment_number} installment sanction letter for project {proposal.unique_id}.")

            messages.success(request, 'Sanction Letter submitted successfully.')

            return redirect('admin:index')
    else:
        form = SanctionLetterForm()

    return render(request, 'admin/submit_sanction_letter.html', {'form': form, 'proposal': proposal})

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def submit_inspection_letter(request, proposal_id):
    proposal = get_object_or_404(Proposal, unique_id=proposal_id)

    if request.method == 'POST':
        form = InspectionReportForm(request.POST, request.FILES)
        if form.is_valid():
            # Save the form to create a new InspectionReport instance
            inspection_report = form.save(commit=False)
            inspection_report.proposal = proposal
            inspection_report.save()

            # Send email to user
            send_inspection_report_email(proposal, inspection_report)

            # Create notification
            create_notification(proposal.user, f"Please view your inspection report for project {proposal.unique_id}.")

            messages.success(request, 'Inspection Letter submitted successfully.')

            return redirect('admin:index')
    else:
        form = InspectionReportForm()

    return render(request, 'admin/submit_inspection_letter.html', {'form': form, 'proposal': proposal})

def send_sanction_letter_email(proposal, sanction_letter,installment_number):
    subject = f"installment sanction letter for project {proposal.unique_id}"
    message = render_to_string('email/sanction_letter_email.html', {'proposal': proposal,'installment_number':installment_number})
    # plain_message = strip_tags(message)
    from_email = EMAIL_HOST_USER  # Set your email address
    to_email = [proposal.user.email]

    email = EmailMessage(subject, message, from_email, to_email)
    email.content_subtype = 'html'
    email.attach_file(sanction_letter.sanction_letter.path)
    email.send()

def send_inspection_report_email(proposal, inspection_report):
    subject = f"Inspection report for project {proposal.unique_id}"
    message = render_to_string('email/inspection_report_email.html', {'proposal': proposal})
    # plain_message = strip_tags(message)
    from_email = EMAIL_HOST_USER  # Set your email address
    to_email = [proposal.user.email]

    email = EmailMessage(subject, message, from_email, to_email)
    
    email.attach_file(inspection_report.inspection_letter.path)
    email.content_subtype = 'html'
    email.send()

def create_notification(user, message):
    created_at=timezone.now()
    Notification.objects.create(user=user, message=message,created_at=created_at)

#generate progress report 
from io import BytesIO
from django.http import HttpResponse
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.shortcuts import render
from django import forms
from .models import ProgressReportDocument, Proposal
from docx import Document
from django.conf import settings
from urllib.parse import quote

def generate_progress_report_document(proposal_unique_id, form_data):
    proposal = Proposal.objects.get(unique_id=proposal_unique_id)
    document = Document()
    document.add_heading(f'Revolving Fund Progress Report - Proposal ID: {proposal}', level=1)

    # # Add form data to the document
    # for field_name, field_value in form_data.items():
    #     document.add_paragraph(f"{field_name}: {field_value}")
    
    # Add form data to the document
    # Add form data to the document
    for field_name, field_value in form_data.items():
        if hasattr(field_value, 'file') and callable(getattr(field_value, 'file', None)):
            document_path = quote(f'media/documents/{field_name}')
            document_url = f'{settings.MEDIA_URL}{document_path}'

        # Add the clickable link to the document
            document.add_paragraph(f"{field_name}: <a href='{document_url}'>{field_name}</a>", style='Hyperlink')
        else:
            document.add_paragraph(f"{field_name}: {field_value}")
    
    

    # Save the document content to a BytesIO object
    temp_file = BytesIO()
    document.save(temp_file)

    # Create a ContentFile from the BytesIO content
    content_file = ContentFile(temp_file.getvalue())

    # Save the document to the database
    report_document = ProgressReportDocument(
        proposal_unique_id=proposal,
        quarter=form_data['quarter'],
        financial_year=form_data['financial_year'],
    )
    report_document.document.save(f'{proposal_unique_id}_progress_report.docx', content_file)
    report_document.save()

    # Clean up the BytesIO object
    temp_file.close()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=revolving_fund_progress_report_{proposal_unique_id}.docx'
    document.save(response)

    return response

from django.shortcuts import render
from .forms import SummaryReportForm
from django.template.loader import get_template

SUBCOMPONENT_CHOICES = [
    ("WMS_RevolvingFund", "WMS: 1.Creation of Revolving Fund for Marketing of Raw Wool"),
    ("EPortal", "WMS: 2.E-Portal for Marketing Auction of Wool and Development of MIS"),
    ("WMS_SelfHelpGroup", "WMS: 3.Financial Assistance for Formation of Wool Producers Societies/Self Help Group(SHGs)"),
    ("WMS_BuyerSellerExpo", "WMS: 4.Organizing Buyers Sellers Meets"),
    ("WMS_InfrastructureDevelopment", "WMS: 5.Financial Assistance to Strengthening Infrastructure Required for Wool Marketing"),
    ("WoolenExpo", "WMS: 6.Organization of Domestic Independent Woolen Expo"),
    ("WoolenExpoHiring", "WMS: 7.Organizing Domestic Expo on Hiring Stall Basis"),
    ("WPS_CFC", "WPS: 1.Establishing Common Facility Centres (CFCs) for Wool Processing Machines/Facilities"),
    ("WPS_SheepShearingMaching", "WPS: 2.Financial Assistance for Sheep Shearing Machines"),
    ("WPS_Equipment", "WPS: 3.Financial Assistance for Other Machines and Equipments"),
    ("WPSSmallToolsDistribution", "WPS: 4.Financial Assistance for Distribution of Small Tools for Manufacturing of Woolen Items"),
    ("HRD_ShortTermProgramme", "HRD: 1.Short Term Training Program for Manufacturing and Weaving of Woolen Items"),
    ("HRD_OnsiteTraining", "HRD: 2.On-Site Training for Industrial Workers"),
    ("HRD_ShearingMachineTraining", "HRD: 3.Training on Machine Sheep Shearing"),
    ("RD", "HRD: 4.Research and Development Projects"),
    ("DomesticMeeting", "HRD: 5.International/Domestic Corporations Stakeholders Meeting/Conference"),
    ("OrganisingSeminar", "HRD: 6.Organizing Seminars, Workshops, Sheep Mela, Fare, Meet"),
    ("WoolSurvey", "HRD: 7.Wool Survey and Study on Wool Sector"),
    ("WoolTestingLab", "HRD: 8.Operating Existing Wool Testing Lab at Bikaner Including Upgradation and WDTC/ISC at Kullu"),
    ("PublicityMonitoring", "HRD: 9.Publicity of Scheme, Monitoring of Projects, Common Visits, Evaluation of Projects/Schemes, and Awareness Program for Swachhta, etc."),
    ("PWDS_PashminaRevolvingFund", "PWDS: 1.Revolving fund for pashmina wool marketing (For UT of J&K & UT of Ladakh)"),
    ("PWDS_PashminaCFC", "PWDS: 2.Setting of machines for pashmina wool processing"),
    ("ShelterShedConstruction", "PWDS: 3.Construction of shelter shed with guard rooms for pashmina goat"),
    ("PortableTentDist", "PWDS: 4.Distribution of portable tents with accessories"),
    ("PredatorProofLightsDist", "PWDS: 5.Distribution of predator-proof corral with LED lights"),
    ("TestingEquipment", "PWDS: 6.Testing equipment, including DNA analyzer for identification/testing of pashmina products"),
    ("ShowroomDevelopment", "PWDS: 7.Development of showroom at Dehairing Plant premises at Leh"),
    ("FodderLandDevelopment", "PWDS: 8.Development of fodder land/Govt. farms for pashmina goats")
]

from .forms import generate_financial_years

FINANCIAL_YEAR_CHOICES = [
    ('', 'All Financial Years'),
    *generate_financial_years(),
]

from django.http import HttpResponse
from docx import Document
from io import BytesIO
from datetime import datetime
from .models import WPS_SheepShearingMaching, WPS_Equipment, WPSSmallToolsDistribution, Proposal

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def summ_report(request, proposal_id):
    proposal = get_object_or_404(Proposal, unique_id=proposal_id)

    if request.method == 'POST':
        form = SummaryReportForm(request.POST)
        if form.is_valid():
            print(form.cleaned_data)
            selected_scheme = form.cleaned_data['scheme']
            selected_subcomponents = form.cleaned_data['subcomponent']
            selected_quarters = []
            if form.cleaned_data['quarter'][0] == '':
                selected_quarters = ['Q1', 'Q2', 'Q3', 'Q4']
            else:
                selected_quarters = form.cleaned_data['quarter']
            selected_years = [] 
            if form.cleaned_data['financial_year'][0] == '':
                selected_years = [choice[0] for choice in FINANCIAL_YEAR_CHOICES[1:]]  
            else:
                selected_years = form.cleaned_data['financial_year']

            # print(selected_years, selected_quarters)
            doc = Document()
            doc.add_heading('Summary Report', level=1)
            ts = datetime.now()
            doc.add_paragraph(f'Generated at: {datetime.now()}' + '\n' + f'Scheme: {", ".join(selected_scheme) if "" not in selected_scheme else "WMS, WPS, PWDS, HRD"}' + '\n' + f'Subcomponents: {"All subcomponents" if "" in selected_subcomponents else selected_subcomponents.join(", ")}')
            
            if 'WMS' in selected_scheme or selected_scheme[0]=='':
                # Check if selected subcomponents start with "WMS"
                matching_subcomponents = []
                if selected_subcomponents[0] == '':
                    matching_subcomponents =  [subcomponent for subcomponent, description in SUBCOMPONENT_CHOICES if description.startswith("WMS")]
                else:
                    for subcomponent in selected_subcomponents:
                        if subcomponent.startswith("WMS"):
                            matching_subcomponents += [sc for sc, description in SUBCOMPONENT_CHOICES if description == subcomponent]
                        else:
                            break

                if matching_subcomponents:
                    summ_reports = []
                    # print(matching_subcomponents)

                    doc.add_heading('\n' + 'Wool Processing Scheme (WMS)', level=2)

                    for subcomp in matching_subcomponents:
                        # print(subcomp)
                        doc.add_heading(f'{subcomp}', level=3)
                        # Match class name with the string sent and get the reports
                        # print(globals()[subcomp].objects.all())
                        summ_reports = list(globals()[subcomp].objects.filter(quarter__in = selected_quarters, financial_year__in = selected_years))

                        for report in summ_reports:
                            pr = Proposal.objects.filter(unique_id = report.proposal_unique_id)
                            goals = pr.values_list('goals', flat=True)[0]  # Assuming there's only one proposal
                            goals_list = [goal for sublist in goals for goal in sublist.values()]
                            num_goals = len(goals_list)
                            num_completed_goals = sum(1 for goal in goals_list if goal['completed'])

                            doc.add_paragraph(f'----------------------------------------------------------------------------------------------------------------------' + '\n' + 
                                              f'Component: {pr.values_list("scheme_component", flat=True)[0]}' + '\n' 
                                              f'Proposal Unique ID: {report.proposal_unique_id}' + '\n' + 
                                              f'Quarter: {report.quarter}' + ", " + f'{report.financial_year}' + '\n' + 
                                              f'Quarterly Allocated Budget: {report.quarterly_allocated_budget}' + '\n' + 
                                              f'Total Quarterly Budget Spent: {report.total_quarterly_budget_spent}')
                            
                            # Write goals in list bullets format
                            doc.add_paragraph('Goals:')
                            for goal in goals_list:
                                goal_completed = "Completed" if goal["completed"] else "Not Completed"
                                doc.add_paragraph(f'- {goal["text"]} ({goal_completed})', style='ListBullet')

                            # Calculate project progress
                            project_progress = num_completed_goals / num_goals if num_goals != 0 else 0
                            doc.add_paragraph(f'Project Progress: {project_progress:.2%}')
                        # summ_reports = []
                                       
            if "WPS" in selected_scheme or selected_scheme[0]=='':
                # Check if selected subcomponents start with "WMS"
                matching_subcomponents = []
                if selected_subcomponents[0] == '':
                    # print([subcomponent for subcomponent, description in SUBCOMPONENT_CHOICES])
                    matching_subcomponents =  [subcomponent for subcomponent, description in SUBCOMPONENT_CHOICES if description.startswith("WPS")]
                else:
                    for subcomponent in selected_subcomponents:
                        if subcomponent.startswith("WPS"):
                            matching_subcomponents += [sc for sc, description in SUBCOMPONENT_CHOICES if description == subcomponent]
                        else:
                            break
                            
                # print(matching_subcomponents)
                if matching_subcomponents:
                    summ_reports = []
                    # print(matching_subcomponents)

                    doc.add_heading('\n' + 'Wool Processing Scheme (WPS)', level=2)

                    summ_reports = list(globals()[subcomp].objects.filter(quarter__in = selected_quarters, financial_year__in = selected_years))

                    for report in summ_reports:
                        pr = Proposal.objects.filter(unique_id = report.proposal_unique_id)
                        goals = pr.values_list('goals', flat=True)[0]  # Assuming there's only one proposal
                        goals_list = [goal for sublist in goals for goal in sublist.values()]
                        num_goals = len(goals_list)
                        num_completed_goals = sum(1 for goal in goals_list if goal['completed'])

                        doc.add_paragraph(f'----------------------------------------------------------------------------------------------------------------------' + '\n' + 
                                            f'Component: {pr.values_list("scheme_component", flat=True)[0]}' + '\n' 
                                            f'Proposal Unique ID: {report.proposal_unique_id}' + '\n' + 
                                            f'Quarter: {report.quarter}' + ", " + f'{report.financial_year}' + '\n' + 
                                            f'Quarterly Allocated Budget: {report.quarterly_allocated_budget}' + '\n' + 
                                            f'Total Quarterly Budget Spent: {report.total_quarterly_budget_spent}')
                        
                        # Write goals in list bullets format
                        doc.add_paragraph('Goals:')
                        for goal in goals_list:
                            goal_completed = "Completed" if goal["completed"] else "Not Completed"
                            doc.add_paragraph(f'- {goal["text"]} ({goal_completed})', style='ListBullet')

                        # Calculate project progress
                        project_progress = num_completed_goals / num_goals if num_goals != 0 else 0
                        doc.add_paragraph(f'Project Progress: {project_progress:.2%}')


            if "HRD" in selected_scheme or selected_scheme[0]=='':
                # Check if selected subcomponents start with "WMS"
                matching_subcomponents = []
                if selected_subcomponents[0] == '':
                    matching_subcomponents =  [subcomponent for subcomponent, description in SUBCOMPONENT_CHOICES if description.startswith("HRD")]
                else:
                    for subcomponent in selected_subcomponents:
                        if subcomponent.startswith("HRD"):
                            matching_subcomponents += [sc for sc, description in SUBCOMPONENT_CHOICES if description == subcomponent]
                        else:
                            break

                if matching_subcomponents:
                    summ_reports = []
                    # print(matching_subcomponents)

                    doc.add_heading('\n' + 'HRD', level=2)

                    summ_reports = list(globals()[subcomp].objects.filter(quarter__in = selected_quarters, financial_year__in = selected_years))

                    for report in summ_reports:
                        pr = Proposal.objects.filter(unique_id = report.proposal_unique_id)
                        goals = pr.values_list('goals', flat=True)[0]  # Assuming there's only one proposal
                        goals_list = [goal for sublist in goals for goal in sublist.values()]
                        num_goals = len(goals_list)
                        num_completed_goals = sum(1 for goal in goals_list if goal['completed'])

                        doc.add_paragraph(f'----------------------------------------------------------------------------------------------------------------------' + '\n' + 
                                            f'Component: {pr.values_list("scheme_component", flat=True)[0]}' + '\n' 
                                            f'Proposal Unique ID: {report.proposal_unique_id}' + '\n' + 
                                            f'Quarter: {report.quarter}' + ", " + f'{report.financial_year}' + '\n' + 
                                            f'Quarterly Allocated Budget: {report.quarterly_allocated_budget}' + '\n' + 
                                            f'Total Quarterly Budget Spent: {report.total_quarterly_budget_spent}')
                        
                        # Write goals in list bullets format
                        doc.add_paragraph('Goals:')
                        for goal in goals_list:
                            goal_completed = "Completed" if goal["completed"] else "Not Completed"
                            doc.add_paragraph(f'- {goal["text"]} ({goal_completed})', style='ListBullet')

                        # Calculate project progress
                        project_progress = num_completed_goals / num_goals if num_goals != 0 else 0
                        doc.add_paragraph(f'Project Progress: {project_progress:.2%}')


            if "PWDS" in selected_scheme or selected_scheme[0]=='':
                # Check if selected subcomponents start with "WMS"
                matching_subcomponents = []
                if selected_subcomponents[0] == '':
                    matching_subcomponents =  [subcomponent for subcomponent, description in SUBCOMPONENT_CHOICES if description.startswith("PWDS")]
                else:
                    for subcomponent in selected_subcomponents:
                        if subcomponent.startswith("PWDS"):
                            matching_subcomponents += [sc for sc, description in SUBCOMPONENT_CHOICES if description == subcomponent]
                        else:
                            break
                    
                # print(matching_subcomponents)
                if matching_subcomponents:
                    summ_reports = []

                    doc.add_heading('\n' + 'Pashmina Wool Development Scheme (PWDS)', level=2)

                    for subcomp in matching_subcomponents:
                        doc.add_heading(f'{subcomp}', level=3)
                        # Match class name with the string sent and get the reports
                        summ_reports = list(globals()[subcomp].objects.filter(quarter__in = selected_quarters, financial_year__in = selected_years))

                        for report in summ_reports:
                            pr = Proposal.objects.filter(unique_id = report.proposal_unique_id)
                            goals = pr.values_list('goals', flat=True)[0]  # Assuming there's only one proposal
                            goals_list = [goal for sublist in goals for goal in sublist.values()]
                            num_goals = len(goals_list)
                            num_completed_goals = sum(1 for goal in goals_list if goal['completed'])

                            doc.add_paragraph(f'----------------------------------------------------------------------------------------------------------------------' + '\n' + 
                                              f'Component: {pr.values_list("scheme_component", flat=True)[0]}' + '\n' 
                                              f'Proposal Unique ID: {report.proposal_unique_id}' + '\n' + 
                                              f'Quarter: {report.quarter}' + ", " + f'{report.financial_year}' + '\n' + 
                                              f'Quarterly Allocated Budget: {report.quarterly_allocated_budget}' + '\n' + 
                                              f'Total Quarterly Budget Spent: {report.total_quarterly_budget_spent}')
                            
                            # Write goals in list bullets format
                            doc.add_paragraph('Goals:')
                            for goal in goals_list:
                                goal_completed = "Completed" if goal["completed"] else "Not Completed"
                                doc.add_paragraph(f'- {goal["text"]} ({goal_completed})', style='ListBullet')

                            # Calculate project progress
                            project_progress = num_completed_goals / num_goals if num_goals != 0 else 0
                            doc.add_paragraph(f'Project Progress: {project_progress:.2%}')
                                                    
            # Save DOCX to BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            # Create response with DOCX file
            response = HttpResponse(doc_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="Admin Custom Summary Report:{ts}.docx"'
            return response

    else:
        form = SummaryReportForm()

    proposals = Proposal.objects.filter(status="Approved").values_list('unique_id', flat=True).distinct()
    proposals = list(proposals)
    return render(request, 'admin/summ_report.html', {'form': form, 'proposals': proposals})

def process(df, gender=0, category=0, state=0, beneficiaries=0):
    """takes all xlsx data and flags of data required
    to analyze and return as just 1 row"""

    analysis = {"state": {}}

    if state:
        all_states_and_ut = ['Andhra Pradesh', 'Arunachal Pradesh', 'Assam', 'Bihar', 'Chhattisgarh', 'Goa', 'Gujarat', 'Haryana', 'Himachal Pradesh', 'Jharkhand', 'Karnataka', 'Kerala', 'Madhya Pradesh', 'Maharashtra', 'Manipur', 'Meghalaya', 'Mizoram', 'Nagaland', 'Odisha', 'Punjab', 'Rajasthan', 'Sikkim', 'Tamil Nadu', 'Telangana', 'Tripura', 'Uttar Pradesh', 'Uttarakhand', 'West Bengal', 'Andaman and Nicobar Islands', 'Chandigarh', 'Dadra and Nagar Haveli and Daman and Diu', 'Lakshadweep', 'Delhi', 'Puducherry']
        
        for state_name in all_states_and_ut:
            state_df = df[df['State/UT (no abbreviations, proper capitalization, max 30 characters)'] == state_name]
            state_analysis = {}

            if gender:
                gender_counts = state_df['Gender (Male/Female/Other)'].value_counts()
                state_analysis['gender'] = [gender_counts.get('Male', 0), gender_counts.get('Female', 0), gender_counts.get('Other', 0)]

            if category:
                category_counts = state_df['Category (General/SC/ST/BPL/OBC)'].value_counts()
                state_analysis['category'] = [category_counts.get('General', 0), category_counts.get('SC', 0), category_counts.get('ST', 0), 
                                              category_counts.get('BPL', 0), category_counts.get('OBC', 0)]

            if beneficiaries:
                state_analysis['beneficiaries'] = len(state_df)

            analysis['state'][state_name] = state_analysis

    return analysis

from .models import BeneficiaryData, ExpenditureData
import json

def revolving_fund_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WMSRevolvingFundForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
                

            print(form.cleaned_data)
            df = pd.read_excel(form.cleaned_data['wool_procured_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # print(form.cleaned_data['proposal_unique_id'], len(df), form.cleaned_data['quarter'], form.cleaned_data['financial_year']) 
            # # #scheme bhi bhejo idhar
            # print(form.cleaned_data['total_quarterly_budget_spent'])
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            df = pd.read_excel(form.cleaned_data['wool_sold_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            
            
            # print(form.cleaned_data['proposal_unique_id'], len(df), form.cleaned_data['quarter'], form.cleaned_data['financial_year']) 
            # # #scheme bhi bhejo idhar
            # print(form.cleaned_data['total_quarterly_budget_spent'])
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()    

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            # Generate and save progress report document
            generate_progress_report_document(proposal_unique_id, form.cleaned_data)
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WMSRevolvingFundForm(initial=initial_data)

    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
        
    return render(request, 'progressReports/WMS/1.RevolvingFund.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from .models import EPortal
from .forms import EPortalForm
from django.http import HttpResponse

from django.shortcuts import render
from django.http import HttpResponse
from .forms import EPortalForm
from .models import EPortal,ExpenditureData
import pandas as pd 

def eportal_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = EPortalForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()            
            
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            generate_progress_report_document(proposal_unique_id, form.cleaned_data)
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = EPortalForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WMS/2.EPortal.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

def send_progress_report_reminders():
    proposals = Proposal.objects.filter(status='Approved')
    # print("done")
    for proposal in proposals:
        current_quarter = get_current_quarter()
        financial_year = get_financial_year()
        
        # Skip if reminder_quarter and reminder_financial_year are None
        if proposal.reminder_quarter is None and proposal.reminder_financial_year is None:
            proposal.reminder_quarter = current_quarter
            proposal.reminder_financial_year = financial_year
            proposal.save()
            continue
        
        if current_quarter != proposal.reminder_quarter or financial_year != proposal.reminder_financial_year:
            message = f'Reminder: Submit progress report for Project ID {proposal.unique_id} of Quarter {proposal.reminder_quarter} ({proposal.reminder_financial_year}).'
            notification = Notification(user=proposal.user, message=message)
            notification.save()

            subject = 'Progress Report Reminder'
            email_template = 'email/progress_report_email_template.html'
            email_content = render_to_string(email_template, {'user': proposal.user, 'proposal': proposal})
            sender_email = settings.EMAIL_HOST_USER
            recipient_email = proposal.user.email
            send_mail(subject, '', sender_email, [recipient_email], html_message=email_content)

            proposal.reminder_quarter = current_quarter
            proposal.reminder_financial_year = financial_year
            proposal.save()


import os
import json
from datetime import datetime, timedelta
from django.apps import apps
from django.core.serializers import serialize

def dump_folderwise_backup():
    try:
        # Get the current datetime
        current_datetime = datetime.now()

        # Get the datetime of 7 days ago
        past_datetime = current_datetime - timedelta(days=7)

        # Backup data of all models modified in the past 7 days
        for model in apps.get_models():
            if model._meta.abstract:
                continue

            if 'updated_at' in [field.name for field in model._meta.fields]:
                queryset = model.objects.filter(updated_at__gte=past_datetime)
            else:
                queryset = model.objects.all()

            if queryset.exists():
                # Create a folder for the model if it doesn't exist
                model_folder = os.path.join("backup", "folder-wise-backup", model._meta.object_name)
                if not os.path.exists(model_folder):
                    os.makedirs(model_folder)

                # Serialize and save backup data for each instance
                for instance in queryset:
                    instance_data = serialize('json', [instance])
                    instance_data = json.loads(instance_data)[0]
                    instance_id = instance_data['pk']
                    instance_filename = f"{instance_id}_{current_datetime.strftime('%Y-%m-%d_%H-%M-%S')}.json"
                    instance_path = os.path.join(model_folder, instance_filename)
                    with open(instance_path, 'w') as f:
                        json.dump(instance_data['fields'], f)

        print("Folderwise backup done.")

    except Exception as e:
        # Handle exceptions such as network errors, timeouts, etc.
        raise RuntimeError(f'Error dumping backup: {str(e)}')

def dump_whole_backup():
    try:
        # Get the current datetime
        current_datetime = datetime.now()

        # Get the datetime of 7 days ago
        past_datetime = current_datetime - timedelta(days=7)

        # Format the datetime as a string
        # datetime_str = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

        # Define the filename with the datetime
        filename = f"backup.json"

        # Define the backup folder
        backup_folder = os.path.join("backup", "whole-backup")

        # Create the backup folder if it doesn't exist
        if not os.path.exists(backup_folder):
            os.makedirs(backup_folder)

        # Define backup path
        backup_path = os.path.join(backup_folder, filename)

        # Backup data of all models modified in the past 7 days
        backup_data = {}
        for model in apps.get_models():
            if model._meta.abstract:
                continue

            # print(model._meta.object_name)
            # try:
            if 'updated_at' in [field.name for field in model._meta.fields]:
                queryset = model.objects.filter(updated_at__gte=past_datetime)
                # prin
            # except AttributeError:
            else:
                # If the model doesn't have an 'updated_at' attribute, backup all data
                queryset = model.objects.all()

            if queryset.exists():
                # Print model fields as a table
                print(f"Model: {model._meta.object_name}")
                print("Field Name".ljust(30), "Type".ljust(20), "Nullable")
                print("-" * 70)
                for field in model._meta.fields:
                    print(field.name.ljust(30), str(field.__class__.__name__).ljust(20), str(field.null).ljust(8))
                print()

                # Print model content
                print("Model Content:")
                for obj in queryset:
                    print(obj.__dict__)
                    print()

                # # there's an issue with using TabularInline outside of the Django admin interface.
                # inline_admin = TabularInline(model, model._admin_site)
                # inline_admin.model = model
                # inline_admin.get_queryset(request=None)
                # print(inline_admin.get_fields(request=None))

                data = json.loads(serialize('json', queryset))
                backup_data[model._meta.object_name] = data

        # Save backup data to a JSON file
        with open(backup_path, 'w') as f:
            json.dump(backup_data, f)

        print(f"Whole Backup successfully created: {backup_path}")

    except Exception as e:
        # Handle exceptions such as network errors, timeouts, etc.
        raise RuntimeError(f'Error dumping backup: {str(e)}')



import json
from datetime import datetime, timedelta
from django.apps import apps
from django.core import serializers
import requests


def take_backup():
    try:
        # Define the time range for backup (past 7 days)
        end_date = datetime.now()
        start_date = end_date - timedelta(days=7)

        # Get all models from installed apps
        models = apps.get_models()

        # List to store serialized data
        backup_data = []

        # Loop through each model
        for model in models:
            if model._meta.abstract:
                continue  # Skip abstract models

            # Check if the model has 'updated_at' field (assuming it's a timestamp of last modification)
            if 'updated_at' in [field.name for field in model._meta.fields]:
                # Fetch all objects updated in the last 7 days
                updated_objects = model.objects.filter(updated_at__gte=start_date, updated_at__lte=end_date)

                # Serialize the queryset and add it to the backup data
                serialized_objects = serializers.serialize('json', updated_objects)
                backup_data.append(serialized_objects)

        # Merge all serialized data into a single JSON string
        backup_json = '[' + ','.join(backup_data) + ']'

        # Specify the target IP address and the endpoint on the Ubuntu VM
        # target_ip = '192.168.1.57:5000'  # Replace with the actual IP address of your Ubuntu VM
        target_ip = '127.0.0.1:5000'  # Replace with the actual IP address of your Ubuntu VM
        endpoint = '/receive_backup'   # Replace with the actual endpoint on your Ubuntu VM

        # Construct the URL
        url = f'http://{target_ip}{endpoint}'
        cert_path = 'D:/IITJodhpur/MTP/CWDB/pems/cert.pem'  # Path to your SSL certificate
        key_path = 'D:/IITJodhpur/MTP/CWDB/pems/key.pem'    # Path to your private key
        # passphrase = read_passphrase_from_config()
        # print(passphrase)

        # Prepare the data payload
        data = {
            'backup_data': backup_json
        }

        # Convert data to JSON
        payload = json.dumps(data)
        bearer_token = 0
        bearer_token_file = "bearer_token.txt"
        with open(bearer_token_file, "r") as file:
            bearer_token = file.read().strip()

        # Example of authentication headers (replace with actual authentication method)
        headers = {
            'Authorization': f'Bearer {bearer_token}',  # Replace with actual access token or authentication method
            'Content-Type': 'application/json'
        }

        # # Create a custom SSL context to specify the SSL/TLS version and use the passphrase
        # ssl_context = create_urllib3_context()
        # ssl_context.load_cert_chain(cert_path, key_path, password=passphrase)

        # # Create a PoolManager with the custom SSL context
        # http = urllib3.PoolManager(cert_reqs='CERT_REQUIRED', ca_certs=cert_path, ssl_context=ssl_context)

        # # Send the POST request securely
        # response = http.request('POST', url, body=payload.encode('utf-8'), headers=headers, timeout=10)
        # response = requests.post(url, data=payload, headers=headers, verify=cert_path, cert=(cert_path, key_path), timeout=10, allow_redirects=True)
        response = requests.post(url, data=payload, headers=headers, timeout=10)
        # print(response)
        if response.status_code == 200:
            return 'Backup sent successfully'
        else:
            response_json = response.json()
            error_message = response_json.get('message') or response_json.get('error') or response.content
            raise RuntimeError(f'Error sending backup: {error_message}')

    except Exception as e:
        # Handle exceptions such as network errors, timeouts, etc.
        raise RuntimeError(f'Error during backup: {str(e)}')

from django.core.mail import send_mail
from django.template.loader import render_to_string

@login_required
def progress_report(request):
    user = request.user
    proposals = Proposal.objects.filter(user=user, status='Approved')
    
    # send_progress_report_reminders()
            
    
    return render(request, 'progressReports/progressreport.html', {'proposals': proposals})


from django.shortcuts import render, redirect
from .models import WMS_SelfHelpGroup, BeneficiaryData, ExpenditureData
from .forms import WMSSelfHelpGroupForm
from django.http import HttpResponse

def selfhelpgroup_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WMSSelfHelpGroupForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            print(form.cleaned_data['quarterly_allocated_budget'], form.cleaned_data['total_quarterly_budget_spent'])
            df = pd.read_excel(form.cleaned_data['shg_members_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # print(form.cleaned_data['proposal_unique_id'], len(df), form.cleaned_data['quarter'], form.cleaned_data['financial_year']) 
            # # #scheme bhi bhejo idhar
            # print(form.cleaned_data['total_quarterly_budget_spent'])
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WMSSelfHelpGroupForm(initial=initial_data)
    
    # Calculate the current quarter based on the created_at date of the Proposal
    # current_quarster = None
    # proposals = Proposal.objects.all()  # Retrieve all proposals, you might need to filter this queryset
    # for proposal in proposals:
    #     created_at = proposal.created_at
    #     if created_at:
    #         current_year = created_at.year
    #         month = created_at.month
    #         if month in range(1, 4):
    #             current_quarter = 'Q4'  # January-March
    #         elif month in range(4, 7):
    #             current_quarter = 'Q1'  # April-June
    #         elif month in range(7, 10):
    #             current_quarter = 'Q2'  # July-September
    #         else:
    #             current_quarter = 'Q3'  # October-December
    #         break  # Found the current quarter, exit loop    
    
    # context = {'form': form}
    # Retrieve goals for the current quarter
    # current_quarter_goals = None
    # print(current_quarter)
    # if current_quarter:
    #     # Assuming goals are stored in the Proposal instance as a list of lists
    #     current_quarter_goals = proposals.filter(created_at__month__in=(1, 4, 7, 10), created_at__year=current_year).values_list('goals', flat=True)
    #     print(current_quarter_goals)
        
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WMS/3.SelfHelpGroup.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

    # return render(request, 'progressReports/WMS/3.SelfHelpGroup.html', {'form': form, 'current_quarter': current_quarter, 'current_quarter_goals': current_quarter_goals})
from django.shortcuts import render, redirect
from .models import WMS_BuyerSellerExpo, ExpenditureData, BeneficiaryData
from .forms import WMSBuyerSellerExpoForm
from django.http import HttpResponse

def buyersellerexpo_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WMSBuyerSellerExpoForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['wool_sellers_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WMSBuyerSellerExpoForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WMS/4.BuyerSellerExpo.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import WMS_InfrastructureDevelopment, ExpenditureData
from .forms import WMSInfrastructureDevelopmentForm
from django.http import HttpResponse

def infrastructuredevelopment_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WMSInfrastructureDevelopmentForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WMSInfrastructureDevelopmentForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WMS/5.infrastructureDevelopment.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from .models import WoolenExpo, WoolenExpoHiring, ExpenditureData
from .forms import WoolenExpoForm, WoolenExpoHiringForm

def woolen_expo(request,proposal_unique_id):
    if request.method == 'POST':
        form = WoolenExpoForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['stall_allotees_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WoolenExpoForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WMS/6.WoolenExpo.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

def woolen_expo_hiring(request,proposal_unique_id):
    if request.method == 'POST':
        form = WoolenExpoHiringForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['stall_allotees_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WoolenExpoHiringForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WMS/7.WoolenExpoHiring.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

# views.py
from django.shortcuts import render, redirect
from .models import WPS_CFC, ExpenditureData
from .forms import WPSCFCForm
from django.http import HttpResponse

def cfc_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WPSCFCForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['facility_user_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WPSCFCForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WPS/1.CFC.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from .forms import WPSSheepShearingMachingForm
from .models import ExpenditureData
def sheep_shearing_machining(request, proposal_unique_id):
    if request.method == 'POST':
        form = WPSSheepShearingMachingForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['beneficiaries_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            # print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
                quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WPSSheepShearingMachingForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WPS/2.SheepShearingMaching.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from .forms import WPSEquipmentForm
from .models import ExpenditureData
def equipment(request, proposal_unique_id):
    if request.method == 'POST':
        form = WPSEquipmentForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['beneficiaries_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WPSEquipmentForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WPS/3.Equipment.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render
from django.http import HttpResponse
from .forms import WPSSmallToolsDistributionForm
from .models import ExpenditureData

def small_tools_distribution(request, proposal_unique_id):
    if request.method == 'POST':
        form = WPSSmallToolsDistributionForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['beneficiaries_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WPSSmallToolsDistributionForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/WPS/4.SmallToolsDistribution.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from .models import HRD_ShortTermProgramme, ExpenditureData
from .forms import HRDShortTermProgrammeForm

def short_term_programme(request, proposal_unique_id):
    if request.method == 'POST':
        form = HRDShortTermProgrammeForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['trainee_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            df = pd.read_excel(form.cleaned_data['master_trainer_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            # print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            df = pd.read_excel(form.cleaned_data['office_assistant_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()                    

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   
            
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = HRDShortTermProgrammeForm(initial=initial_data)

    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/1.ShortTermProgramme.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from .models import HRD_OnsiteTraining, ExpenditureData
from .forms import HRDOnsiteTrainingForm

def onsite_training_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = HRDOnsiteTrainingForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['trainee_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()    

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = HRDOnsiteTrainingForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/2.OnsiteTraining.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import HRD_ShearingMachineTraining
from .forms import HRDShearingMachineTrainingForm
from django.http import HttpResponse

def shearing_machine_training_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = HRDShearingMachineTrainingForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['trainee_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = HRDShearingMachineTrainingForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/3.ShearingMachineTraining.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from .models import RD
from .forms import RDForm

def rd_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = RDForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = RDForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/4.R&D.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import DomesticMeeting, ExpenditureData
from .forms import DomesticMeetingForm
from django.http import HttpResponse

def domestic_meeting_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = DomesticMeetingForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['participants_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = DomesticMeetingForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/5.DomesticMeeting.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import OrganisingSeminar
from .forms import OrganisingSeminarForm
from django.http import HttpResponse

def organising_seminar_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = OrganisingSeminarForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['participants_details_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = OrganisingSeminarForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/6.OrganisingSeminar.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import WoolSurvey, ExpenditureData
from .forms import WoolSurveyForm
from django.http import HttpResponse

def wool_survey_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WoolSurveyForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WoolSurveyForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/7.WoolSurvey.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import WoolTestingLab, ExpenditureData
from .forms import WoolTestingLabForm
from django.http import HttpResponse

def wool_testing_lab_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WoolTestingLabForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['details_of_trainees_wdtc_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WoolTestingLabForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/8.WoolTestingLab.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import PublicityMonitoring, ExpenditureData
from .forms import PublicityMonitoringForm
from django.http import HttpResponse

def publicity_monitoring_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PublicityMonitoringForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = PublicityMonitoringForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/HRD/9.PublicityMonitoring.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

# pwds
from django.shortcuts import render, redirect
from .models import PWDS_PashminaRevolvingFund, BeneficiaryData, ExpenditureData
from .forms import PWDS_PashminaRevolvingFundForm
from django.http import HttpResponse


def pashmina_revolving_fund_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PWDS_PashminaRevolvingFundForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['wool_procured_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # print(form.cleaned_data['proposal_unique_id'], len(df), form.cleaned_data['quarter'], form.cleaned_data['financial_year']) 
            # # #scheme bhi bhejo idhar
            # print(form.cleaned_data['total_quarterly_budget_spent'])
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            df = pd.read_excel(form.cleaned_data['wool_sold_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # print(form.cleaned_data['proposal_unique_id'], len(df), form.cleaned_data['quarter'], form.cleaned_data['financial_year']) 
            # # #scheme bhi bhejo idhar
            # print(form.cleaned_data['total_quarterly_budget_spent'])
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()    

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = PWDS_PashminaRevolvingFundForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/PWDS/1.PashminaRevolvingFund.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import PWDS_PashminaCFC, ExpenditureData
from .forms import PWDS_PashminaCFCForm
from django.http import HttpResponse

def pashmina_cfc_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PWDS_PashminaCFCForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Assuming get_financial_year is defined elsewhere
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = PWDS_PashminaCFCForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/PWDS/2.PashminaCFC.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import ShelterShedConstruction, ExpenditureData
from .forms import ShelterShedConstructionForm
from django.http import HttpResponse

def shelter_shed_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = ShelterShedConstructionForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['physical_financial_progress_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # print(form.cleaned_data['proposal_unique_id'], len(df), form.cleaned_data['quarter'], form.cleaned_data['financial_year']) 
            # # #scheme bhi bhejo idhar
            # print(form.cleaned_data['total_quarterly_budget_spent'])
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = ShelterShedConstructionForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/PWDS/3.ShelterShedConstruction.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import PortableTentDist, ExpenditureData
from .forms import PortableTentDistForm
from django.http import HttpResponse

def portable_tent_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PortableTentDistForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['physical_financial_progress_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # print(form.cleaned_data['proposal_unique_id'], len(df), form.cleaned_data['quarter'], form.cleaned_data['financial_year']) 
            # # #scheme bhi bhejo idhar
            # print(form.cleaned_data['total_quarterly_budget_spent'])
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = PortableTentDistForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/PWDS/4.PortableTentDist.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import PredatorProofLightsDist, ExpenditureData
from .forms import PredatorProofLightsDistForm
from django.http import HttpResponse

def predator_proof_lights_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PredatorProofLightsDistForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            df = pd.read_excel(form.cleaned_data['physical_financial_progress_sheet'])
            analysis = process(df, 1, 1, 1, 1)
            print(analysis)
            
            # print(form.cleaned_data['proposal_unique_id'], len(df), form.cleaned_data['quarter'], form.cleaned_data['financial_year']) 
            # # #scheme bhi bhejo idhar
            # print(form.cleaned_data['total_quarterly_budget_spent'])
            
            # Iterate through the state-wise analysis data
            for state_name, state_data in analysis['state'].items():
                # Check if beneficiaries count for the state is greater than 0
                if state_data.get('beneficiaries', 0) > 0:
                    # Create a new instance of BeneficiaryData model
                    beneficiary_data_instance = BeneficiaryData(
                        proposal_unique_id=proposal,
                        num_beneficiaries=state_data.get('beneficiaries', 0),
                        num_general_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[0],
                        num_obc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[4],
                        num_sc_beneficiaries=state_data.get('category', [0, 0, 0, 0, 0])[1] ,
                        num_st_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[2],
                        num_bpl_beneficiaries= state_data.get('category', [0, 0, 0, 0, 0])[3],
                        state_of_beneficiaries=state_name,
                        num_males=state_data.get('gender', [0, 0, 0])[0],
                        num_females=state_data.get('gender', [0, 0, 0])[1],
                        num_other_gender=state_data.get('gender', [0, 0, 0])[2],
                        quarter=form.cleaned_data['quarter'],
                        year=form.cleaned_data['financial_year'],
                        scheme=scheme,
                    )

                    # Save the instance to the database
                    beneficiary_data_instance.save()

            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = PredatorProofLightsDistForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/PWDS/5.PredatorProofLightsDist.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import TestingEquipment, ExpenditureData
from .forms import TestingEquipmentForm
from django.http import HttpResponse

def testing_equipment_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = TestingEquipmentForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = TestingEquipmentForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/PWDS/6.TestingEquipment.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import ShowroomDevelopment, ExpenditureData
from .forms import ShowroomDevelopmentForm
from django.http import HttpResponse

def showroom_development_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = ShowroomDevelopmentForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = ShowroomDevelopmentForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/PWDS/7.ShowroomDevelopment.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})

from django.shortcuts import render, redirect
from .models import FodderLandDevelopment, ExpenditureData
from .forms import FodderLandDevelopmentForm
from django.http import HttpResponse

def fodder_land_development_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = FodderLandDevelopmentForm(request.POST, request.FILES)
        proposal = Proposal.objects.get(unique_id=proposal_unique_id)
        scheme=proposal.project_scheme
        if form.is_valid():
            checkbox_values = request.POST.getlist('checkbox')
            print(checkbox_values)

            # Process the checkbox values as needed
            changes = 0
            for di in checkbox_values:
                changes = 1
                goal_edit = json.loads(di)
                proposal.goals[goal_edit["quarter"]]["goal_"+ str(goal_edit["number"]+1)]["completed"] = 1
            # proposal.goals = [{"goal_1": {"text": "a ceq", "completed": 0}, "goal_2": {"text": "cewc ewcw", "completed": 0}}, {"goal_1": {"text": "cewf ", "completed": 0}, "goal_2": {"text": "cewce ewd", "completed": 0}}]
            if changes:
                proposal.save()
            expenditure_data_instance = ExpenditureData(
                proposal_unique_id=proposal,
                                #added allotted and spent here 
               quarterly_budget_spent = form.cleaned_data['total_quarterly_budget_spent'],
                quarterly_budget_allocated = form.cleaned_data['quarterly_allocated_budget'],
                quarter=form.cleaned_data['quarter'],
                year=form.cleaned_data['financial_year'],
                scheme=scheme,
            )

            # Save the instance to the database
            expenditure_data_instance.save()   

            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponseRedirect(reverse('authapp:form_submitted_successfully'))
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = FodderLandDevelopmentForm(initial=initial_data)
    
    proposal = Proposal.objects.filter(unique_id=proposal_unique_id)
    return render(request, 'progressReports/PWDS/8.FodderLandDevelopmen.html', {'form': form, 'goals': json.dumps(list(proposal.values_list('goals', flat=True))), 'approved_at': proposal.first().approved_at})


#staff view
from django.contrib.auth.views import LoginView
from django.urls import reverse_lazy
class CustomLoginView(LoginView):
    def get_success_url(self):
        user = self.request.user

        if user.is_superuser or user.is_staff:
            return reverse_lazy('admin:index')  # Use reverse_lazy to avoid URL resolution issues
        return reverse_lazy('authapp:dashboard') 
    
    def form_invalid(self, form):
        messages.error(self.request, 'Invalid email or password. Please try again.')
        return super().form_invalid(form)
    
@user_passes_test(lambda u: u.is_staff or u.is_superuser, login_url='/login/')
def staff_dashboard(request):
    # Count the number of proposals for each status
    pending_count = Proposal.objects.filter(status='Pending').count()
    approved_count = Proposal.objects.filter(status='Approved').count()
    completed_count = Proposal.objects.filter(status='Completed').count()
    rejected_count = Proposal.objects.filter(status='Rejected').count()

    # Your view logic goes here

    return render(request, 'staff_template/staff_dashboard.html', {
        'pending_count': pending_count,
        'approved_count': approved_count,
        'completed_count': completed_count,
        'rejected_count': rejected_count,
    })


@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def all_proposal_details(request):
    proposals = Proposal.objects.all()
    return render(request, 'staff_template/admin_proposal_list.html', {'proposals': proposals})

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def submit_approval(request, proposal_id):
    proposal = get_object_or_404(Proposal, unique_id=proposal_id)

    if request.method == 'POST':
        form = ProposalApprovalForm(request.POST, request.FILES,instance=proposal)
        if form.is_valid():
            # Save the form to update the model
            form.save()
            
            # Save the current quarter and set the reminder for the next quarter
            proposal.reminder_quarter = get_current_quarter()
            proposal.reminder_financial_year=get_financial_year()
            proposal.reminder_sent = False
            proposal.approved_at=timezone.now()
            proposal.save()
            # Notify the user about the status change
            send_status_change_notification(proposal.user, proposal.unique_id, proposal.status,proposal.project_sanction_letter)

            messages.success(request, 'Proposal status changed successfully.')
            
            return redirect('admin:index')
    else:
        form = ProposalApprovalForm(instance=proposal)
    return render(request, 'admin/submit_approval.html', {'form': form, 'proposal': proposal})




#projects HomePage 
from django.shortcuts import render
from .models import Proposal
from .forms import ProposalFilterForm

def proposal_list(request):
    form = ProposalFilterForm(request.GET)
    
    if form.is_valid():
        proposals = Proposal.objects.all()
        

    # Filter proposals with status 'Approved' or 'Completed'
        status=form.cleaned_data['status']
        scheme = form.cleaned_data['scheme']
        state= form.cleaned_data['state']
        quarter = form.cleaned_data['quarter']
        financial_year = form.cleaned_data['financial_year']
        
         # Filter proposals with status 'Approved' or 'Completed'
        proposals = Proposal.objects.filter(status__in=['Approved', 'Completed'])

        
        if status:
             proposals = proposals.filter(status=status)
        if scheme:
            proposals = proposals.filter(project_scheme=scheme)
        if state:
            proposals = proposals.filter(implementingAgencyState=state)
            
       # Filter expenditure data based on quarter and financial year
        expenditure_data = ExpenditureData.objects.values('proposal_unique_id').annotate(
            total_budget_allocated=Sum('quarterly_budget_allocated'),
            total_budget_spent=Sum('quarterly_budget_spent')
        )

        if quarter != '':  # Check if a specific quarter is selected
            expenditure_data = expenditure_data.filter(quarter=quarter)

        if financial_year != '':  # Check if a specific financial year is selected
            expenditure_data = expenditure_data.filter(year=financial_year)

        context = {
        'proposals': proposals,
        'form': form,
        'expenditure_data': expenditure_data,
       
    }

        return render(request, 'main/HomePage/Projects.html', context)

from django.shortcuts import render
from django.db.models import Sum
from .models import BeneficiaryData
from .forms import BeneficiaryDataFilterForm

def beneficiary_data_table(request):
    form = BeneficiaryDataFilterForm(request.GET)

    # Filter queryset based on form data
    queryset = BeneficiaryData.objects.all()

    if form.is_valid():
        state = form.cleaned_data.get('state')
        quarter = form.cleaned_data['quarter']
        financial_year = form.cleaned_data.get('financial_year')
        scheme = form.cleaned_data.get('scheme')

        if state:
            queryset = queryset.filter(state_of_beneficiaries=state)
        if quarter:
            queryset = queryset.filter(quarter=quarter)
        if financial_year:
            queryset = queryset.filter(year=financial_year)
        if scheme:
            queryset = queryset.filter(scheme=scheme)

    # Calculate sum for each column
    total_beneficiaries = queryset.aggregate(Sum('num_beneficiaries'))['num_beneficiaries__sum'] or 0
    total_general = queryset.aggregate(Sum('num_general_beneficiaries'))['num_general_beneficiaries__sum'] or 0
    total_obc = queryset.aggregate(Sum('num_obc_beneficiaries'))['num_obc_beneficiaries__sum'] or 0
    total_sc = queryset.aggregate(Sum('num_sc_beneficiaries'))['num_sc_beneficiaries__sum'] or 0
    total_st = queryset.aggregate(Sum('num_st_beneficiaries'))['num_st_beneficiaries__sum'] or 0
    total_bpl = queryset.aggregate(Sum('num_bpl_beneficiaries'))['num_bpl_beneficiaries__sum'] or 0
    total_males = queryset.aggregate(Sum('num_males'))['num_males__sum'] or 0
    total_females = queryset.aggregate(Sum('num_females'))['num_females__sum'] or 0
    total_other_gender = queryset.aggregate(Sum('num_other_gender'))['num_other_gender__sum'] or 0

    context = {
        'form': form,
        'data_table': queryset,
        'total_beneficiaries': total_beneficiaries,
        'total_general': total_general,
        'total_obc': total_obc,
        'total_sc': total_sc,
        'total_st': total_st,
        'total_bpl': total_bpl,
        'total_males': total_males,
        'total_females': total_females,
        'total_other_gender': total_other_gender,
    }

    return render(request, 'main/HomePage/beneficairy_data.html', context)

from django.db import models

from django.shortcuts import render
from .models import FundDistribution, ExpenditureData
from cwdb_admin.models import AdministrativeExpenditure
from .forms import scheme_filterform
from django.db.models import Sum

from django.db.models import Sum

def get_admin_expense_for_year(year):
    admin_expenses = (
        AdministrativeExpenditure.objects
        .filter(financial_year=year)
        .aggregate(total_expense=Sum('admin_exp'))
    )
    return admin_expenses['total_expense'] or 0



def iwdp_view(request):
    form = scheme_filterform(request.GET)
    if form.is_valid():
        financial_year = form.cleaned_data.get('financial_year')
        fund_type=form.cleaned_data.get('select_type')
        fund_allocated=FundDistribution.objects.values('wms', 'wps', 'hrdpa', 'pwds', 'admin_exp', 'financial_year')
        if financial_year:
            fund_allocated=FundDistribution.objects.filter(financial_year=financial_year).values('wms', 'wps', 'hrdpa', 'pwds', 'admin_exp', 'financial_year')
        
        if fund_type=='Fund Allocated':
            fund_data=fund_allocated
            for data in fund_data:
                data['iwdp'] = data['wms'] + data['wps'] + data['hrdpa'] + data['pwds'] + data['admin_exp']

            
        elif fund_type=='Expenditure':
            expenditure_data = ExpenditureData.objects.all()
            if financial_year:
                expenditure_data = expenditure_data.filter(year=financial_year)

            expenditure_data_grouped = expenditure_data.values('year').annotate(
            wms_expenditure=Sum('quarterly_budget_spent', filter=models.Q(scheme='WMS')),
            wps_expenditure=Sum('quarterly_budget_spent', filter=models.Q(scheme='WPS')),
            hrdpa_expenditure=Sum('quarterly_budget_spent', filter=models.Q(scheme='HRDPA')),
            pwds_expenditure=Sum('quarterly_budget_spent', filter=models.Q(scheme='PWDS')),
        )

        # Calculate IWDP expenditure for each year
            for entry in expenditure_data_grouped:
                entry['iwdp_expenditure'] = (
                entry['wms_expenditure'] or 0) + (
                entry['wps_expenditure'] or 0) + (
                entry['hrdpa_expenditure'] or 0) + (
                entry['pwds_expenditure'] or 0)
                
        # Initialize data list
            fund_data = []
            for entry in expenditure_data_grouped:
                admin_exp=get_admin_expense_for_year(entry['year'])
                data_entry = {
                'wms': entry['wms_expenditure'] or 0,
                'wps': entry['wps_expenditure'] or 0,
                'hrdpa': entry['hrdpa_expenditure'] or 0,
                'pwds': entry['pwds_expenditure'] or 0,
                'admin_exp': admin_exp,  # Fetch admin_exp from FundDistribution model, 
                'financial_year': entry['year'],
                'iwdp': entry['iwdp_expenditure']+admin_exp,#also 
            }
                fund_data.append(data_entry)
        elif fund_type=='Fund Sanctioned':
            expenditure_data = ExpenditureData.objects.all()
            if financial_year:
                expenditure_data = expenditure_data.filter(year=financial_year)

            expenditure_data_grouped = expenditure_data.values('year').annotate(
            wms_expenditure=Sum('quarterly_budget_allocated', filter=models.Q(scheme='WMS')),
            wps_expenditure=Sum('quarterly_budget_allocated', filter=models.Q(scheme='WPS')),
            hrdpa_expenditure=Sum('quarterly_budget_allocated', filter=models.Q(scheme='HRDPA')),
            pwds_expenditure=Sum('quarterly_budget_allocated', filter=models.Q(scheme='PWDS')),
        )

        # Calculate IWDP expenditure for each year
            for entry in expenditure_data_grouped:
                entry['iwdp_expenditure'] = (
                entry['wms_expenditure'] or 0) + (
                entry['wps_expenditure'] or 0) + (
                entry['hrdpa_expenditure'] or 0) + (
                entry['pwds_expenditure'] or 0)
                
        # Initialize data list
            fund_data = []
            for entry in expenditure_data_grouped:
                admin_exp = FundDistribution.objects.filter(financial_year=entry['year']).values('admin_exp').first()
                data_entry = {
                'wms': entry['wms_expenditure'] or 0,
                'wps': entry['wps_expenditure'] or 0,
                'hrdpa': entry['hrdpa_expenditure'] or 0,
                'pwds': entry['pwds_expenditure'] or 0,
                'admin_exp': admin_exp['admin_exp'] if admin_exp else 0,  # Fetch admin_exp from FundDistribution model, 
                'financial_year': entry['year'],
                'iwdp': entry['iwdp_expenditure']+admin_exp['admin_exp'] or 0,#also add admin exp
            }
                fund_data.append(data_entry)
            
        else:
            fund_data = FundDistribution.objects.none()
            
        


        
        context = {
                'form':form,
                'financial_year': financial_year,
                'fund_data':fund_data,
               
                
            }
            
        return render(request, 'main/HomePage/iwdp.html', context)
    

from django.shortcuts import render
from .models import FundDistribution, ExpenditureData
from .forms import scheme_filterform
from django.db.models import Sum

def wms_scheme_view(request):
    form = scheme_filterform(request.GET)
    if form.is_valid():
        financial_year = form.cleaned_data.get('financial_year')

        # Fetch Fund Allocated data for WMS scheme
        fund_allocated=FundDistribution.objects.filter().values('wms', 'financial_year')
        if financial_year:
            fund_allocated = FundDistribution.objects.filter(financial_year=financial_year).values('wms', 'financial_year')

        # Fetch Expenditure data for WMS scheme
        expenditure_data = ExpenditureData.objects.filter(scheme='WMS')
        if financial_year:
            expenditure_data = expenditure_data.filter(year=financial_year)

        expenditure_data_grouped = expenditure_data.values('year').annotate(
            wms_expenditure=Sum('quarterly_budget_spent'),
            wms_sanctioned=Sum('quarterly_budget_allocated'),
        )
        
        # Initialize data list
        wms_data = []
        for entry in expenditure_data_grouped:
            # Fetch Fund Allocated value for WMS scheme of that year
            fund_allocated_value = fund_allocated.filter(financial_year=entry['year']).first()
            data_entry = {
                'wms_sanctioned': entry['wms_sanctioned'] or 0,
                'wms_expenditure': entry['wms_expenditure'] or 0,
                'wms_allocated': fund_allocated_value['wms'] if fund_allocated_value else 0,
                'financial_year': entry['year'],
            }
            wms_data.append(data_entry)
        
        # Find financial years in FundDistribution but not in expenditure_data_grouped
        missing_years = set(entry['financial_year'] for entry in fund_allocated) - set(entry['year'] for entry in expenditure_data_grouped)

        # Add entries for missing financial years
        for year in missing_years:
            fund_allocated_value = fund_allocated.filter(financial_year=year).first()
            data_entry = {
                'wms_sanctioned': 0,
                'wms_expenditure': 0,
                'wms_allocated': fund_allocated_value['wms'] if fund_allocated_value else 0,
                'financial_year': year,
            }
            wms_data.append(data_entry)

        context = {
            'form': form,
            'financial_year': financial_year,
            'wms_data': wms_data,
        }

        return render(request, 'main/HomePage/wms_scheme.html', context)


def wps_scheme_view(request):
    form = scheme_filterform(request.GET)
    if form.is_valid():
        financial_year = form.cleaned_data.get('financial_year')

        # Fetch Fund Allocated data for WPS scheme
        fund_allocated = FundDistribution.objects.filter().values('wps', 'financial_year')
        if financial_year:
            fund_allocated = FundDistribution.objects.filter(financial_year=financial_year).values('wps', 'financial_year')

        # Fetch Expenditure data for WPS scheme
        expenditure_data = ExpenditureData.objects.filter(scheme='WPS')
        if financial_year:
            expenditure_data = expenditure_data.filter(year=financial_year)

        expenditure_data_grouped = expenditure_data.values('year').annotate(
            wps_expenditure=Sum('quarterly_budget_spent'),
            wps_sanctioned=Sum('quarterly_budget_allocated'),
        )
        
        # Initialize data list
        wps_data = []
        for entry in expenditure_data_grouped:
            # Fetch Fund Allocated value for WPS scheme of that year
            fund_allocated_value = fund_allocated.filter(financial_year=entry['year']).first()
            data_entry = {
                'wps_sanctioned': entry['wps_sanctioned'] or 0,
                'wps_expenditure': entry['wps_expenditure'] or 0,
                'wps_allocated': fund_allocated_value['wps'] if fund_allocated_value else 0,
                'financial_year': entry['year'],
            }
            wps_data.append(data_entry)
        
        # Find financial years in FundDistribution but not in expenditure_data_grouped
        missing_years = set(entry['financial_year'] for entry in fund_allocated) - set(entry['year'] for entry in expenditure_data_grouped)

        # Add entries for missing financial years
        for year in missing_years:
            fund_allocated_value = fund_allocated.filter(financial_year=year).first()
            data_entry = {
                'wps_sanctioned': 0,
                'wps_expenditure': 0,
                'wps_allocated': fund_allocated_value['wps'] if fund_allocated_value else 0,
                'financial_year': year,
            }
            wps_data.append(data_entry)

        context = {
            'form': form,
            'financial_year': financial_year,
            'wps_data': wps_data,
        }

        return render(request, 'main/HomePage/wps_scheme.html', context)

def pwds_scheme_view(request):
    form = scheme_filterform(request.GET)
    if form.is_valid():
        financial_year = form.cleaned_data.get('financial_year')

        # Fetch Fund Allocated data for PWDs scheme
        fund_allocated = FundDistribution.objects.filter().values('pwds', 'financial_year')
        if financial_year:
            fund_allocated = FundDistribution.objects.filter(financial_year=financial_year).values('pwds', 'financial_year')

        # Fetch Expenditure data for PWDs scheme
        expenditure_data = ExpenditureData.objects.filter(scheme='PWDS')
        if financial_year:
            expenditure_data = expenditure_data.filter(year=financial_year)

        expenditure_data_grouped = expenditure_data.values('year').annotate(
            pwds_expenditure=Sum('quarterly_budget_spent'),
            pwds_sanctioned=Sum('quarterly_budget_allocated'),
        )
        
        # Initialize data list
        pwds_data = []
        for entry in expenditure_data_grouped:
            # Fetch Fund Allocated value for PWDs scheme of that year
            fund_allocated_value = fund_allocated.filter(financial_year=entry['year']).first()
            data_entry = {
                'pwds_sanctioned': entry['pwds_sanctioned'] or 0,
                'pwds_expenditure': entry['pwds_expenditure'] or 0,
                'pwds_allocated': fund_allocated_value['pwds'] if fund_allocated_value else 0,
                'financial_year': entry['year'],
            }
            pwds_data.append(data_entry)
        
        # Find financial years in FundDistribution but not in expenditure_data_grouped
        missing_years = set(entry['financial_year'] for entry in fund_allocated) - set(entry['year'] for entry in expenditure_data_grouped)

        # Add entries for missing financial years
        for year in missing_years:
            fund_allocated_value = fund_allocated.filter(financial_year=year).first()
            data_entry = {
                'pwds_sanctioned': 0,
                'pwds_expenditure': 0,
                'pwds_allocated': fund_allocated_value['pwds'] if fund_allocated_value else 0,
                'financial_year': year,
            }
            pwds_data.append(data_entry)

        context = {
            'form': form,
            'financial_year': financial_year,
            'pwds_data': pwds_data,
        }

        return render(request, 'main/HomePage/pwds_scheme.html', context)


def hrdpa_scheme_view(request):
    form = scheme_filterform(request.GET)
    if form.is_valid():
        financial_year = form.cleaned_data.get('financial_year')

        # Fetch Fund Allocated data for HRDPA scheme
        fund_allocated = FundDistribution.objects.filter().values('hrdpa', 'financial_year')
        if financial_year:
            fund_allocated = FundDistribution.objects.filter(financial_year=financial_year).values('hrdpa', 'financial_year')

        # Fetch Expenditure data for HRDPA scheme
        expenditure_data = ExpenditureData.objects.filter(scheme='HRD')
        if financial_year:
            expenditure_data = expenditure_data.filter(year=financial_year)

        expenditure_data_grouped = expenditure_data.values('year').annotate(
            hrdpa_expenditure=Sum('quarterly_budget_spent'),
            hrdpa_sanctioned=Sum('quarterly_budget_allocated'),
        )
        
        # Initialize data list
        hrdpa_data = []
        for entry in expenditure_data_grouped:
            # Fetch Fund Allocated value for HRDPA scheme of that year
            fund_allocated_value = fund_allocated.filter(financial_year=entry['year']).first()
            data_entry = {
                'hrdpa_sanctioned': entry['hrdpa_sanctioned'] or 0,
                'hrdpa_expenditure': entry['hrdpa_expenditure'] or 0,
                'hrdpa_allocated': fund_allocated_value['hrdpa'] if fund_allocated_value else 0,
                'financial_year': entry['year'],
            }
            hrdpa_data.append(data_entry)
        
        # Find financial years in FundDistribution but not in expenditure_data_grouped
        missing_years = set(entry['financial_year'] for entry in fund_allocated) - set(entry['year'] for entry in expenditure_data_grouped)

        # Add entries for missing financial years
        for year in missing_years:
            fund_allocated_value = fund_allocated.filter(financial_year=year).first()
            data_entry = {
                'hrdpa_sanctioned': 0,
                'hrdpa_expenditure': 0,
                'hrdpa_allocated': fund_allocated_value['hrdpa'] if fund_allocated_value else 0,
                'financial_year': year,
            }
            hrdpa_data.append(data_entry)

        context = {
            'form': form,
            'financial_year': financial_year,
            'hrdpa_data': hrdpa_data,
        }

        return render(request, 'main/HomePage/hrdpa_scheme.html', context)
    
def admin_exp_view(request):
    form = scheme_filterform(request.GET)
    
    if form.is_valid():
        financial_year = form.cleaned_data.get('financial_year')
        fund_type=form.cleaned_data.get('select_type')
        fund_allocated=FundDistribution.objects.values('wms', 'wps', 'hrdpa', 'pwds', 'admin_exp', 'financial_year')
        if financial_year:
            fund_allocated=FundDistribution.objects.filter(financial_year=financial_year).values('wms', 'wps', 'hrdpa', 'pwds', 'admin_exp', 'financial_year')
        
        if fund_type=='Fund Allocated' or fund_type=='Fund Sanctioned':
            fund_data=fund_allocated
            for data in fund_data:
                data['iwdp'] = data['wms'] + data['wps'] + data['hrdpa'] + data['pwds'] + data['admin_exp']

            
        elif fund_type=='Expenditure':
            expenditure_data = ExpenditureData.objects.all()
            if financial_year:
                expenditure_data = expenditure_data.filter(year=financial_year)

            expenditure_data_grouped = expenditure_data.values('year').annotate(
            wms_expenditure=Sum('quarterly_budget_spent', filter=models.Q(scheme='WMS')),
            wps_expenditure=Sum('quarterly_budget_spent', filter=models.Q(scheme='WPS')),
            hrdpa_expenditure=Sum('quarterly_budget_spent', filter=models.Q(scheme='HRDPA')),
            pwds_expenditure=Sum('quarterly_budget_spent', filter=models.Q(scheme='PWDS')),
        )

        # Calculate IWDP expenditure for each year
            for entry in expenditure_data_grouped:
                entry['iwdp_expenditure'] = (
                entry['wms_expenditure'] or 0) + (
                entry['wps_expenditure'] or 0) + (
                entry['hrdpa_expenditure'] or 0) + (
                entry['pwds_expenditure'] or 0)
                
        # Initialize data list
            fund_data = []
            for entry in expenditure_data_grouped:
                admin_exp=get_admin_expense_for_year(entry['year'])
                data_entry = {
                'wms': entry['wms_expenditure'] or 0,
                'wps': entry['wps_expenditure'] or 0,
                'hrdpa': entry['hrdpa_expenditure'] or 0,
                'pwds': entry['pwds_expenditure'] or 0,
                'admin_exp': admin_exp,  # Fetch admin_exp from FundDistribution model, 
                'financial_year': entry['year'],
                'iwdp': entry['iwdp_expenditure']+admin_exp,#also 
            }
                fund_data.append(data_entry)
        
        else:
            fund_data = FundDistribution.objects.none()

        
        context = {
                'form1':form,
                'financial_year': financial_year,
                'fund_data':fund_data,
               
                
            }

        return render(request, 'main/HomePage/admin_exp.html', context)
    
from api.serializers import ProposalSerializer
from django.http import JsonResponse
from .models import Proposal

def proposal_api(request):
    proposal = Proposal.objects.all()
    serializer = ProposalSerializer(proposal, many = True)
    return JsonResponse(serializer.data, safe=False)


from django.shortcuts import render
from .models import Proposal  # Import the Proposal model

def proposal_documents(request):
    # Query all proposals
    user=request.user
    proposals = Proposal.objects.filter(user=user).exclude(status='Pending')

    # Render the template with the proposals data
    return render(request, 'proposal/proposal_documents.html', {'proposals': proposals})

# documents proposal
from django.shortcuts import render, get_object_or_404
from .models import Proposal, SanctionLetter, InspectionReport

def sanction_letters_view(request, proposal_id):
    proposal = get_object_or_404(Proposal, unique_id=proposal_id)
    sanction_letters = SanctionLetter.objects.filter(proposal=proposal).order_by('installment_number')
    context = {'sanction_letters': sanction_letters, 'proposal': proposal}
    return render(request, 'proposal/sanction_letters.html', context)

def inspection_letters_view(request, proposal_id):
    proposal = get_object_or_404(Proposal, unique_id=proposal_id)
    inspection_reports = InspectionReport.objects.filter(proposal=proposal).order_by('-created_at')
    context = {'inspection_reports': inspection_reports, 'proposal': proposal}
    return render(request, 'proposal/inspection_letters.html', context)


from cwdb_admin.models import Index_Notification 

def index_notification_view(request):
    notifications = Index_Notification.objects.order_by('-created_at')
    return render(request, 'main/HomePage/index_notifications.html', {'notifications': notifications})

# views.py

from django.shortcuts import render
@login_required
def form_submitted_successfully(request):
    return render(request, 'main/form_submitted_sucessfully.html')
